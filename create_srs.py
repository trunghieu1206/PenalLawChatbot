import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Khởi tạo document
doc = docx.Document()

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text):
    doc.add_paragraph(text)

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

# Title Page
add_title("TÀI LIỆU ĐẶC TẢ YÊU CẦU PHẦN MỀM (SRS)")
add_title("Dự án: VNPLaw - Vietnamese Penal Law Chatbot")
add_title("Phiên bản 1.0")
doc.add_page_break()

# MỤC LỤC
add_heading("Mục lục", 1)
add_para("1. Lịch sử thay đổi tài liệu")
add_para("2. Giới thiệu chung")
add_para("3. Tổng quan dự án")
add_para("4. Các yêu cầu về giao tiếp bên ngoài")
add_para("5. Các yêu cầu hệ thống")
add_para("6. Tài liệu tham khảo")
doc.add_page_break()

# 1. Lịch sử thay đổi tài liệu
add_heading("Lịch sử thay đổi tài liệu", 1)
table = doc.add_table(rows=2, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Ngày'
hdr_cells[1].text = 'Phiên bản'
hdr_cells[2].text = 'Mô tả thay đổi'
hdr_cells[3].text = 'Tác giả'
row_cells = table.rows[1].cells
row_cells[0].text = '06/05/2026'
row_cells[1].text = '1.0'
row_cells[2].text = 'Khởi tạo tài liệu đặc tả SRS'
row_cells[3].text = 'Development Team'

# 2. Giới thiệu chung
add_heading("Giới thiệu chung", 1)
add_heading("Mục đích", 2)
add_para("Tài liệu này đặc tả chi tiết các yêu cầu chức năng, phi chức năng, hệ thống và kiến trúc cho dự án VNPLaw (Vietnamese Penal Law Chatbot). Hệ thống hướng tới việc xây dựng một trợ lý AI thông minh trong lĩnh vực pháp luật hình sự Việt Nam, ứng dụng mô hình RAG (Retrieval-Augmented Generation) và LangGraph để cung cấp lập luận pháp lý chính xác, có căn cứ và phù hợp với quy định về thời hiệu (Điều 7 BLHS).")

add_heading("Thuật ngữ & từ viết tắt", 2)
terms = [
    "RAG: Retrieval-Augmented Generation - Công nghệ truy xuất tăng cường sinh ngôn ngữ.",
    "LLM: Large Language Model - Mô hình ngôn ngữ lớn (Gemini 2.5 Flash).",
    "BLHS: Bộ luật Hình sự Việt Nam (các phiên bản 1999, 2009, 2015, 2017, 2025).",
    "JWT: JSON Web Token - Tiêu chuẩn xác thực người dùng.",
    "LoRA: Low-Rank Adaptation - Kỹ thuật tinh chỉnh mô hình nhúng (Embedding)."
]
for term in terms:
    add_bullet(term)

# 3. Tổng quan dự án
add_heading("Tổng quan dự án", 1)
add_heading("Khảo sát hiện trạng", 2)
add_para("Việc tra cứu và đối chiếu các phiên bản pháp luật hình sự thủ công đòi hỏi chuyên môn cao, mất nhiều thời gian và dễ xảy ra sai sót, đặc biệt là khi áp dụng nguyên tắc hồi tố có lợi. Hệ thống VNPLaw ra đời nhằm tự động hóa quá trình ánh xạ sự kiện pháp lý vào điều khoản luật, trích xuất tình tiết giảm nhẹ/tăng nặng, và cung cấp cái nhìn đa chiều (Thẩm phán, Bào chữa, Bị hại).")

add_heading("Tổng quan chức năng hệ thống", 2)
funcs = [
    "Tư vấn pháp lý hình sự dựa trên đầu vào của người dùng (tình tiết vụ án).",
    "Đóng vai trò thiên kiến (Bias mode): Thẩm phán (trung lập), Luật sư bào chữa (giảm nhẹ), Luật sư bị hại (tăng nặng).",
    "Chế độ Luyện tập (Practice mode): Đánh giá, chấm điểm phân tích của người dùng dựa trên barem chuẩn.",
    "Quản lý lịch sử tư vấn (Chat Sessions) bằng tài khoản cá nhân.",
    "Hệ thống quản trị (Admin Dashboard) theo dõi số lượng người dùng (Track-visit), thống kê vụ án và quản lý phản hồi đánh giá."
]
for f in funcs:
    add_bullet(f)

add_heading("Biểu đồ use case tổng quan", 3)
add_para("Hệ thống bao gồm 2 nhóm Tác nhân (Actor) chính: Người dùng (bao gồm Khách và Người dùng đã đăng nhập) và Quản trị viên.")
add_para("- Người dùng: Chat tư vấn pháp lý, Luyện tập phân tích, Đăng nhập/Đăng ký, Xem lịch sử, Phản hồi câu trả lời.")
add_para("- Quản trị viên: Theo dõi Dashboard thống kê, Quản lý danh sách phản hồi từ người dùng.")

add_heading("Biểu đồ use case phân rã Tư vấn pháp lý", 3)
add_para("Trong Use Case tư vấn, quá trình được phân rã thành: Nhập thông tin sự kiện -> Trích xuất dữ kiện (hành vi, ngày phạm tội) -> Tra cứu semantic các điều luật liên quan -> Ánh xạ điều khoản (xử lý hồi tố) -> Sinh câu trả lời theo góc nhìn vai trò.")

add_heading("Quy trình nghiệp vụ (Pipeline LangGraph)", 3)
add_para("Hệ thống AI xử lý theo luồng đồ thị trạng thái (StateGraph):")
add_para("1. classify_intent -> 2. extract_facts -> 3. clarification_check -> 4. multi_query_rewrite -> 5. parallel_retrieve -> 6. temporal_priority_tagger -> 7. rerank -> 8. map_laws -> 9. generate/rebuttal.")

add_heading("Đặc tả chức năng", 2)
add_heading("Đặc tả use case Phân tích vụ án", 3)
add_para("Người dùng nhập đoạn mô tả vụ việc. Hệ thống tự động trích xuất các trường bắt buộc (hành vi, ngày phạm tội). Nếu thiếu, hệ thống yêu cầu bổ sung. Nếu đủ, hệ thống truy xuất Milvus Lite và sinh lập luận pháp lý có kèm bảng 'ĐIỀU KHOẢN ÁP DỤNG'.")

add_heading("Đặc tả chức năng Quản lý Dashboard", 3)
add_para("Quản trị viên truy cập trang Dashboard. Hệ thống hiển thị tổng số phiên trò chuyện, thống kê loại tội phạm, thống kê theo vai trò phân tích, và danh sách User-stats lấy từ PostgreSQL thông qua API Spring Boot.")

# 4. Các yêu cầu về giao tiếp bên ngoài
add_heading("Các yêu cầu về giao tiếp bên ngoài", 1)
add_heading("Giao diện người dùng (User Interface)", 2)
add_para("Hệ thống sử dụng React 19 và Vite, cung cấp giao diện Web Responsive. Sidebar bên trái quản lý lịch sử chat. Sidebar bên phải tra cứu nhanh toàn văn các điều luật. Giao diện Chat chính hiển thị các bong bóng tin nhắn (MessageBubble) hỗ trợ Markdown.")

add_heading("Giao diện phần cứng (Hardware Interface)", 2)
add_para("Hệ thống tối ưu hóa chạy trên CPU Bare-metal (Ubuntu 22.04), sử dụng OMP_NUM_THREADS và MKL_NUM_THREADS để tận dụng tối đa số luồng vật lý của CPU. Tuy nhiên, hệ thống tự động nhận diện và sử dụng CUDA GPU nếu có sẵn.")

add_heading("Giao diện phần mềm (Software Interface)", 2)
add_para("Hệ thống là sự kết hợp của 4 tiến trình độc lập: Nginx (Port 80), Spring Boot Backend (Port 8080), FastAPI AI Service (Port 8000), và PostgreSQL (Port 5432). Giao tiếp nội bộ qua RESTful API.")

add_heading("Giao tiếp với các hệ thống khác (Communication Interface)", 2)
add_para("Giao tiếp với API OpenRouter để sử dụng mô hình LLM google/gemini-2.5-flash. Giao tiếp với kho lưu trữ HuggingFace để tải LoRA Adapter cho mô hình nhúng (Embedding).")

# 5. Các yêu cầu hệ thống
add_heading("Các yêu cầu hệ thống", 1)
add_heading("Yêu cầu chức năng", 2)
add_para("Hệ thống phải đảm bảo các chức năng: Xác thực JWT, Lưu trữ Session DB, Tri-Path Retrieval (Semantic + Pinned), Edition-Aware Law Mapping (Xử lý hồi tố Điều 7), Đánh giá Practice Mode.")

add_heading("Yêu cầu tích hợp", 2)
add_para("Tích hợp thư viện Pymilvus để kết nối CSDL Vector cục bộ (Milvus Lite). Tích hợp thư viện langgraph để quản lý luồng hội thoại trạng thái của AI Agent.")

add_heading("Yêu cầu phi chức năng", 2)
add_heading("Yêu cầu về hiệu năng", 3)
add_para("Quá trình rerank bằng mô hình Cross-Encoder (BAAI/bge-reranker-v2-m3) phải xử lý context 8192 token mượt mà để chứa trọn vẹn văn bản luật dài nhất (Điều 232 BLHS 2017). Thời gian phản hồi toàn bộ luồng RAG phải dưới 30 giây đối với CPU và dưới 10 giây đối với GPU.")

add_heading("Yêu cầu an toàn", 3)
add_para("Tất cả các dữ liệu văn bản pháp lý đẩy vào mô hình AI và CSDL đều phải đi qua hàm sanitize_text() nhằm loại bỏ các ký tự surrogate UTF-16 lỗi, tránh làm treo bộ mã hóa JSON của Python.")

add_heading("Yêu cầu về bảo mật", 3)
add_para("API Back-end được bảo mật bằng Spring Security JWT. Mật khẩu lưu trữ trong CSDL được mã hóa BCrypt. Các API Admin (/api/admin/**) chỉ cho phép truy cập với quyền ROLE_ADMIN.")

add_heading("Các yêu cầu khác", 2)
add_para("Theo dõi người dùng truy cập một cách ẩn danh (Deduplication) sử dụng UUID ở localStorage kết hợp constraint Unique ở Database.")

# 6. Tài liệu tham khảo
add_heading("Tài liệu tham khảo", 1)
ref_list = [
    "rag_architecture_design.md: Kiến trúc chi tiết luồng LangGraph.",
    "toaan_gov_datasets_2.json: Nguồn dữ liệu huấn luyện."
]
for ref in ref_list:
    add_bullet(ref)

# Lưu file
output_path = "VNPLaw_SRS.docx"
doc.save(output_path)
print(f"Successfully generated {output_path}")
