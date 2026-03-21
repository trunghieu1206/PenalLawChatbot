"""
parse_docx.py — Parse Vietnamese penal law DOCX files into structured JSON.

Each output record = one Article, annotated with its Chapter and effective dates.

Usage:
    python scripts/parse_docx.py

Output files (in laws_documents_raw/):
    VB_1999.json, VB_2009.json, VB_2017.json, VB_2025.json

Install dependency (if not already):
    pip install python-docx
"""

import re
import json
from pathlib import Path
from docx import Document

# ─── Configuration ───────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent.parent / "laws_documents_raw"

FILES = [
    {
        "path": BASE_DIR / "VB_1999.docx",
        "source": "BLHS 1999",
        "effective_start": "2000-07-01",
        "effective_end": "2009-12-31",
    },
    {
        "path": BASE_DIR / "VB_2009.docx",
        "source": "BLHS 2009 (sửa đổi)",
        "effective_start": "2010-01-01",
        "effective_end": "2017-12-31",
    },
    {
        "path": BASE_DIR / "VB_2017.docx",
        "source": "BLHS 2015 (sửa đổi 2017)",
        "effective_start": "2018-01-01",
        "effective_end": None,   # None = still in effect
    },
    {
        "path": BASE_DIR / "VB_2025.docx",
        "source": "BLHS 2025",
        "effective_start": "2025-07-01",
        "effective_end": None,
    },
]

# ─── Regex Patterns ───────────────────────────────────────────────────────────
# Matches: "Chương I", "Chương II", "Chương XIV", "Chương 1" etc.
RE_CHAPTER = re.compile(
    r"^chương\s+([IVXLCDM]+|\d+)[\.\s]*[-–—]?\s*(.*)",
    re.IGNORECASE
)

# Matches: "Điều 1.", "Điều 12.", "Điều 12a.", "Điều 12 a."
RE_ARTICLE = re.compile(
    r"^điều\s+(\d+\s*[a-z]?)\s*[.\-–—]?\s*(.*)",
    re.IGNORECASE
)

# ─── Helpers ─────────────────────────────────────────────────────────────────

def clean(text: str) -> str:
    """Normalize whitespace, remove zero-width chars."""
    text = text.replace("\u200b", "").replace("\xa0", " ")
    return " ".join(text.split()).strip()


def extract_paragraphs(doc: Document) -> list[str]:
    """
    Extract all non-empty paragraph texts from a DOCX document.
    Tables are also scanned (some DOCX files put content in table cells).
    """
    paragraphs = []
    for para in doc.paragraphs:
        t = clean(para.text)
        if t:
            paragraphs.append(t)
    # Also scan tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = clean(para.text)
                    if t and t not in paragraphs:
                        paragraphs.append(t)
    return paragraphs


def parse_docx(cfg: dict) -> list[dict]:
    """
    Parse a single DOCX file into a list of article records.

    Each record:
    {
        "article_number":   "51",
        "title":            "Các tình tiết giảm nhẹ trách nhiệm hình sự",
        "chapter_number":   "IV",
        "chapter":          "CÁC BIỆN PHÁP TƯ PHÁP",
        "content":          "1. Các tình tiết sau đây...",
        "source":           "BLHS 2015 (sửa đổi 2017)",
        "effective_start":  "2018-01-01",
        "effective_end":    null
    }
    """
    path: Path = cfg["path"]
    if not path.exists():
        print(f"  ⚠️  File not found: {path}")
        return []

    print(f"\n📖 Parsing: {path.name}")
    doc = Document(str(path))
    paragraphs = extract_paragraphs(doc)
    print(f"  → {len(paragraphs)} paragraphs extracted")

    articles = []
    current_chapter_num = ""
    current_chapter_title = ""
    current_article_num = None
    current_article_title = ""
    content_lines = []

    def flush_article():
        """Save the current article buffer if it has content."""
        nonlocal current_article_num, current_article_title, content_lines
        if current_article_num is not None:
            content = "\n".join(content_lines).strip()
            if content:
                articles.append({
                    "article_number":  current_article_num.strip(),
                    "title":           current_article_title,
                    "chapter":         current_chapter_num,
                    "content":         content,
                    "source":          cfg["source"],
                    "effective_start": cfg["effective_start"],
                    "effective_end":   cfg["effective_end"],
                })
        current_article_num = None
        current_article_title = ""
        content_lines = []

    for line in paragraphs:
        ch_match = RE_CHAPTER.match(line)
        art_match = RE_ARTICLE.match(line)

        if ch_match:
            # Save previous article before switching chapter
            flush_article()
            current_chapter_num = ch_match.group(1).strip()
            current_chapter_title = clean(ch_match.group(2))
            print(f"  📂 Chương {current_chapter_num}: {current_chapter_title[:60]}")

        elif art_match:
            # Save previous article, start new one
            flush_article()
            current_article_num = art_match.group(1).strip()
            current_article_title = clean(art_match.group(2))

        else:
            # Body text — belongs to the current article
            if current_article_num is not None:
                content_lines.append(line)

    # Don't forget the last article
    flush_article()

    print(f"  ✅ {len(articles)} articles parsed")
    return articles


def main():
    total = 0
    for cfg in FILES:
        articles = parse_docx(cfg)
        if not articles:
            continue

        out_path = cfg["path"].with_suffix(".json")
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(articles, f, ensure_ascii=False, indent=2)

        print(f"  💾 Saved → {out_path.name}  ({len(articles)} articles)")
        total += len(articles)

    print(f"\n🎉 Done. Total articles across all files: {total}")
    print("\nNext step:")
    print("  python scripts/ingest_laws.py --file laws_documents_raw/VB_2017.json "
          '--source "BLHS 2015 (sửa đổi 2017)" --date 2018-01-01')


if __name__ == "__main__":
    main()
